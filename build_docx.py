#!/usr/bin/env python3
# -*- coding: utf-8 -*-
import re
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.opc.constants import RELATIONSHIP_TYPE as RT

GREEN = RGBColor(0x2F, 0x5B, 0x4C)
INK = RGBColor(0x22, 0x22, 0x22)
GREY = RGBColor(0x66, 0x66, 0x66)

doc = Document()

sec = doc.sections[0]
sec.page_height = Cm(29.7)
sec.page_width = Cm(21.0)
for m in ("top_margin", "bottom_margin", "left_margin", "right_margin"):
    setattr(sec, m, Cm(0.9))

normal = doc.styles["Normal"]
normal.font.name = "Microsoft YaHei"
normal.font.size = Pt(8.5)
rpr = normal.element.get_or_add_rPr()
rf = rpr.get_or_add_rFonts()
rf.set(qn("w:eastAsia"), "微软雅黑")


def para(text, size=8.5, after=1.5, before=0, line=1.0, base_bold=False,
         color=INK, align=None, indent=None):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_after = Pt(after)
    pf.space_before = Pt(before)
    pf.line_spacing = line
    if align is not None:
        p.alignment = align
    if indent is not None:
        pf.left_indent = Cm(indent)
    for i, seg in enumerate(re.split(r"\*\*", text)):
        if seg == "":
            continue
        r = p.add_run(seg)
        r.font.size = Pt(size)
        r.bold = base_bold or (i % 2 == 1)
        r.font.color.rgb = color
    return p


def heading(text):
    return para(text, size=9.5, before=3, after=1.2, base_bold=True, color=GREEN)


def bullet(text, size=8.5, after=0.8, color=INK):
    return para("• " + text, size=size, after=after, indent=0.4, color=color)


def add_hyperlink(p, url, text):
    part = p.part
    r_id = part.relate_to(url, RT.HYPERLINK, is_external=True)
    h = OxmlElement("w:hyperlink")
    h.set(qn("r:id"), r_id)
    r = OxmlElement("w:r")
    rp = OxmlElement("w:rPr")
    c = OxmlElement("w:color"); c.set(qn("w:val"), "1A5FB4"); rp.append(c)
    u = OxmlElement("w:u"); u.set(qn("w:val"), "single"); rp.append(u)
    sz = OxmlElement("w:sz"); sz.set(qn("w:val"), "19"); rp.append(sz)
    r.append(rp)
    t = OxmlElement("w:t"); t.text = text; r.append(t)
    h.append(r)
    p._p.append(h)


# ================= TITLE =================
para("陪训 PetCoach · AI 宠物行为训练教练", size=13.5, after=0.5, base_bold=True,
     color=GREEN, align=WD_ALIGN_PARAGRAPH.CENTER)
para("方向：宠物 AI　｜　一句话：免费训狗教程铺天盖地，但新手养狗人常在见效前就放弃。"
     "PetCoach 用 AI 把口语描述转成**行为诊断 → 严重度分级 → 分周训练计划 → 每日任务**，"
     "并以打卡 + 规则化次日提示演示“陪主人坚持”的最小闭环——**我们卖的不是知识，是坚持到见效**。",
     size=8.5, after=1.5, align=WD_ALIGN_PARAGRAPH.CENTER, color=GREY)

# ================= ① =================
heading("① 机会判断")
bullet("**目标用户**：一线 / 新一线城市 90、00 后新手养狗人（狗 0–3 岁），独居或上班族，"
       "被拆家、分离焦虑、随地大小便、乱吠、扑人咬人击穿。白皮书报道 90 后占宠主 **42.7%**、00 后 **26.3%**（增速最快）。")
bullet("**核心痛点**：缺的不是方法，是**坚持与反馈**。训练见效常需**数周至数月**，用户却往往几天就放弃；"
       "线下代训 **¥1000–5800/月甚至上万**且“回家就反弹”；免费视频不个性化、无反馈、不纠错。")
bullet("**为什么现在**：① 2025 城镇犬猫消费 **3126 亿**，服务占比仅 **6.5%**（医疗 27.6%），渗透率洼地；"
       "② 付费主力换成习惯 App/AI 订阅打卡的 90/00 后；③ 多模态 LLM 边际成本趋零，未来可“看视频给即时反馈”，"
       "是唯一能做高频微反馈的形态（当前先用文本模型跑通主线）；④ 线下“买证、粗暴训练、代训反弹”信任崩塌，"
       "透明正向强化成为现成卖点。")
bullet("**主动亮风险**：打卡 / 习惯类 App 自身留存也差——用户凭什么对我的 App 坚持数周？"
       "“AI 能否真正提升坚持率”是本方向命门，必须用真实 D14 留存验证，而非靠 Demo 说服。")

# ================= ② =================
heading("② 用户证据（可核查链接 + 社群观察）")
bullet("**痛点高频（公开讨论）**：知乎分离焦虑讨论反复出现“嚎叫、刨门、破坏、独处焦虑”求助（知乎问答 / 专栏）；"
       "权威科普将分离焦虑症状列为吠叫 / 嚎叫、啃咬抓门、踱步（Farmina）。分离焦虑占犬行为问题约 **20–40%**。")
bullet("**成熟付费市场**：羊城晚报报道到府课 **¥1000–3000/节**；CBNData / 市值榜报道基础服从训练普遍 **>¥3000**、"
       "有机构 **¥5800/月**；央视网报道宠物学校 1–2 个月 **¥3500–4500**、线下训练从几千到上万——即使效果不稳定，用户仍在花钱。")
bullet("**小红书 / 淘宝真实评论（社群观察）**：小红书用户“出门练习这个点，要等狗不叫了再回来，"
       "还是就是出门两分钟就回来，不管它叫不叫”（反映方法执行难、缺即时反馈）；淘宝训犬课买家“之前还花了快一万块"
       "把他送去训狗基地，原本有好转了但一出门还是被打回原形，疯狂拆家，有吃的也不吃，已经被投诉好几次了……已经想弃养了”"
       "（线下贵且反弹、痛点极强）。原始截图见代码仓库 evidence/ 目录。")
bullet("**市场盘**：央视网《2026 中国宠物行业白皮书》报道——2025 城镇犬猫消费 **3126 亿**，单犬年均 **¥3006**，"
       "食品 / 医疗 / 用品 / 服务 = 53.7% / 27.6% / 12.2% / **6.5%**。")
bullet("**证据强度自陈**：以上为公开网络资料 / 桌面调研 + 社群观察，**尚未做一手访谈或问卷**；"
       "“3–7 天放弃”是我从求助帖归纳的待验证假设，不是实测数据，已列为 ⑤ 第一步的校准目标。", color=GREY)

# ================= ③ =================
heading("③ MVP 样例（可运行 Demo，非原型图）")
bullet("**三段式 AI 工作流（核心资产）**：Stage1 把零散口语**结构化**成宠物画像；Stage2 输出"
       "**行为根因诊断 + 严重度分级(low/medium/high) + 医疗边界**（疑似疾病即建议就医）；"
       "Stage3 生成 **4–8 周训练计划**（每周目标、每日任务“做什么 / 多久 / 成功标准”、所需道具、安全红线）。"
       "三段 Prompt 均强约束只输出 JSON，前端稳定渲染。")
bullet("**技术栈**：Next.js(App Router) + TypeScript + Tailwind；`/api/coach` 串行编排三段；"
       "DeepSeek live / 确定性 mock 双模式；Zod 输入校验 + live 输出结构兜底（异常回退 mock）；Vitest 测试覆盖。")
bullet("**已完成**：结果页（诊断卡 / 严重度刻度 / 6 周计划手风琴 / 安全红线 / 一次性免责声明）、"
       "本地打卡进度、localStorage 持久化、基于完成率的规则化次日提示（最小陪跑闭环）。")
bullet("**诚实边界**：当前“次日提示”是前端规则演示，**不是 AI 回传微调**；deepseek-chat 不读取图片 / 视频——"
       "这两块（AI 微调 + 视频理解）是下一步、也是真正的护城河。", color=GREY)

# ================= ④ =================
heading("④ 商业判断")
bullet("**谁付费**：愿为狗花钱（单犬年均 ¥3006）、习惯线上订阅、被线下“贵且反弹”劝退的新手养狗人。")
bullet("**为什么付费**：买的不是知识（免费遍地），是“**坚持执行 + 持续反馈 + 真见效**”——比线下便宜约 10 倍、"
       "可迁移到主人、全程透明正向强化。")
bullet("**定价**：免费诊断 + 周计划做**获客钩子** → **¥99「6 周训练营」一次性**（锚定线下 ¥1000+，一次性买断"
       "降低“中途没坚持就退订”的流失，转化优于月订阅） → 训练中推荐围栏 / 漏食球 / 嗅闻垫 / 零食，赚 **CPS 导购佣金**。")
bullet("**获客**：小红书 / 抖音 before-after 打卡视频（天然可传播）、狗主社群裂变、宠粮 / 用品 KOL 合作、"
       "长尾 SEO（“XX 犬拆家怎么办”）。**诚实标注**：付费意愿目前仅间接证据，真实转化用 waitlist + 假付费按钮实测；"
       "before-after 依赖真实见效，不在验证前当作已成立。")

# ================= ⑤ =================
heading("⑤ 2 周验证计划（每步：验证什么 + 资源 + go / no-go）")
bullet("**Week 1 · 需求验证（尚未开始，为计划）**：深访 **10 位**新手狗主（小红书私信 / 社群 / 问卷星），"
       "校准“放弃时点”与付费意愿；上线 Landing Page，投 **¥200–500** 测“点击 → 留资”。"
       "**阈值**：留资率 ≥ 10%；≥ 6/10 能说出明确放弃时点或卡住原因。")
bullet("**Week 2 · 产品验证**：邀 **10 人真实试用 14 天**（窗口按“见效需数周”设 14 天，不是 7 天）。"
       "核心指标 = **D7/D14 留存 + 每日打卡完成率 + 连续 3 天完成占比 + 假付费点击转化 + NPS**。"
       "**go/no-go**：D14 留存 ≥ 40% 且连续 3 天完成率 ≥ 50% 且假付费转化 ≥ 15% → 继续；否则回炉重想陪跑机制。"
       "n=10 只取定性信号与方向，不做统计结论。")
bullet("**资源投入**：个人时间 4–6h/天、AI API 约 ¥数百、Vercel 免费部署、¥200–500 投放、剪辑 / 问卷工具。**无需大预算**。")

# ================= FOOTER =================
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(4)
p.paragraph_format.space_after = Pt(0)
r = p.add_run("在线 Demo："); r.bold = True; r.font.size = Pt(8.5); r.font.color.rgb = INK
add_hyperlink(p, "https://pet-coach-mvp.vercel.app/", "pet-coach-mvp.vercel.app")
r = p.add_run("　｜　代码仓库："); r.bold = True; r.font.size = Pt(8.5); r.font.color.rgb = INK
add_hyperlink(p, "https://github.com/bbainthug/pet-coach-mvp", "github.com/bbainthug/pet-coach-mvp")
r = p.add_run("　｜　本产品定位行为训练辅助，不替代兽医诊疗。")
r.font.size = Pt(8.3); r.font.color.rgb = GREY

doc.save("pet-coach-A4.docx")
print("saved pet-coach-A4.docx")
